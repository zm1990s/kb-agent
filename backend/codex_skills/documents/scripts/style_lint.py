#!/usr/bin/env python3
"""Style lint for DOCX (python-docx).

Goal: quickly surface formatting inconsistencies that typically cause
"why does this paragraph look different" issues:

- direct (run-level) formatting overrides (font, size, bold/italic/color)
- direct paragraph formatting overrides (space before/after, indents)
- heading-like paragraphs that are not actually Heading styles
- font usage summary

This is intentionally conservative: it does not mutate the doc.
"""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document


def _iter_paragraphs(doc: Document):
    # body paragraphs
    for p in doc.paragraphs:
        yield p
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    # headers/footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def _has_direct_run_formatting(run) -> bool:
    f = run.font
    return any(
        v is not None
        for v in [
            run.bold,
            run.italic,
            run.underline,
            f.name,
            f.size,
            f.color.rgb if f.color else None,
        ]
    )


def _has_direct_paragraph_formatting(p) -> bool:
    pf = p.paragraph_format
    return any(
        v is not None
        for v in [
            pf.left_indent,
            pf.right_indent,
            pf.first_line_indent,
            pf.space_before,
            pf.space_after,
            pf.line_spacing,
        ]
    )


def _looks_like_heading(p) -> bool:
    # heuristic: short line, bold-ish, ends without period, no list prefix
    txt = (p.text or "").strip()
    if not txt:
        return False
    if len(txt) > 80:
        return False
    if txt.endswith((".", ";", ":")):
        return False
    if re.match(r"^\d+(?:\.\d+)*\s+", txt):
        # numbered headings are common; allow
        pass
    # any bold run counts
    if any(r.bold for r in p.runs if r.text and r.bold):
        return True
    # style name might already imply
    return False


def main() -> None:
    ap = argparse.ArgumentParser(description="Lint a DOCX for common style/formatting issues")
    ap.add_argument("input_docx")
    ap.add_argument(
        "--json", dest="json_out", default=None, help="Write a JSON report to this path"
    )
    args = ap.parse_args()

    doc = Document(args.input_docx)

    font_names = Counter()
    direct_runs = 0
    direct_paras = 0
    heading_like_not_heading = []

    # collect examples (limit)
    examples = defaultdict(list)

    for i, p in enumerate(_iter_paragraphs(doc), start=1):
        style_name = p.style.name if p.style is not None else ""

        if _has_direct_paragraph_formatting(p):
            direct_paras += 1
            if len(examples["direct_paragraph_formatting"]) < 5:
                examples["direct_paragraph_formatting"].append(
                    {"para_index": i, "text": p.text[:120], "style": style_name}
                )

        if _looks_like_heading(p) and not style_name.startswith("Heading"):
            heading_like_not_heading.append(
                {"para_index": i, "text": p.text[:120], "style": style_name}
            )

        for r in p.runs:
            if r.text:
                if r.font and r.font.name:
                    font_names[r.font.name] += len(r.text)
                if _has_direct_run_formatting(r):
                    direct_runs += 1
                    if len(examples["direct_run_formatting"]) < 5:
                        examples["direct_run_formatting"].append(
                            {
                                "para_index": i,
                                "run_text": r.text[:80],
                                "style": style_name,
                            }
                        )

    report = {
        "input": str(Path(args.input_docx).resolve()),
        "fonts_by_char_count": dict(font_names.most_common()),
        "direct_run_formatting_runs": direct_runs,
        "direct_paragraph_formatting_paragraphs": direct_paras,
        "heading_like_paragraphs_not_heading_style": heading_like_not_heading[:20],
        "examples": dict(examples),
        "notes": [
            "Direct formatting is not always wrong, but it often causes inconsistent output when templates change.",
            "Heading-like paragraphs not using Heading styles can break TOC and accessibility.",
        ],
    }

    if args.json_out:
        Path(args.json_out).write_text(json.dumps(report, indent=2), encoding="utf-8")

    # human-friendly summary
    print("[style_lint] direct run-formatting runs:", direct_runs)
    print("[style_lint] direct paragraph-formatting paragraphs:", direct_paras)
    if font_names:
        top = ", ".join([f"{k}({v})" for k, v in font_names.most_common(5)])
        print("[style_lint] top fonts by char count:", top)
    if heading_like_not_heading:
        print("[style_lint] heading-like paragraphs not using Heading styles (first 10):")
        for item in heading_like_not_heading[:10]:
            print(f"  - #{item['para_index']}: style='{item['style']}' text='{item['text']}'")


if __name__ == "__main__":
    main()
