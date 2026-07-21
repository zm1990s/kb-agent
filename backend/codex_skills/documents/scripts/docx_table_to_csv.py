#!/usr/bin/env python3
"""Export a DOCX table to CSV.

Why this exists
---------------
Round-tripping tables between spreadsheets and DOCX is common in reporting workflows.
This exporter is intentionally simple: it extracts plain cell text.

Usage examples
--------------
python scripts/docx_table_to_csv.py input.docx --table_index 0 --out table0.csv

Notes
-----
- Merged cells are exported as their visible text in each grid cell (Word stores these differently).
- Multi-paragraph cells are joined with newlines.
"""

from __future__ import annotations

import argparse
import csv
import os
import sys
from pathlib import Path

from docx import Document


def main() -> None:
    ap = argparse.ArgumentParser(description="Export a DOCX table to CSV")
    ap.add_argument("docx", type=Path)
    ap.add_argument("--table_index", type=int, default=0)
    ap.add_argument("--out", type=Path, required=True)
    args = ap.parse_args()

    if not args.docx.exists():
        raise FileNotFoundError(args.docx)

    doc = Document(str(args.docx))
    if not doc.tables:
        raise RuntimeError("No tables found in DOCX")
    if args.table_index < 0 or args.table_index >= len(doc.tables):
        raise RuntimeError(f"table_index out of range (0..{len(doc.tables) - 1})")

    t = doc.tables[args.table_index]
    args.out.parent.mkdir(parents=True, exist_ok=True)

    with args.out.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        for row in t.rows:
            out_row = []
            for cell in row.cells:
                parts = []
                for p in cell.paragraphs:
                    txt = (p.text or "").strip("\n")
                    if txt:
                        parts.append(txt)
                out_row.append("\n".join(parts))
            w.writerow(out_row)

    print(f"[OK] Exported table {args.table_index} → {args.out}")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        if os.environ.get("DOCS_DEBUG") == "1":
            raise
        print(f"[ERROR] {e}", file=sys.stderr)
        raise SystemExit(2)
