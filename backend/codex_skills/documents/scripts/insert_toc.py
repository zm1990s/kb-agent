#!/usr/bin/env python3
"""Insert a Table of Contents (TOC) field at a placeholder paragraph.

Default behavior
----------------
- Finds a paragraph whose full text equals the placeholder token (default: "[[TOC]]").
- Replaces it with a Word TOC field code (complex field with fldChar begin/separate/end).
- Sets the document setting "updateFields on open" so Word refreshes TOC/page/refs when opened.

Why this exists
---------------
TOCs are extremely common in reports, and "TOC looks wrong" is often because:
- headings aren't using Heading 1/2/3 styles, or
- fields were not updated before rendering/export.

This script doesn't *compute* the TOC. Word (or a GUI editor) still needs to update fields.
"""

from __future__ import annotations

import argparse
import os
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _set_update_fields_on_open(docx_path: Path) -> None:
    """Ensure word/settings.xml contains <w:updateFields w:val="true"/>."""

    tmp = Path(tempfile.mkdtemp(prefix="docx_toc_patch_"))
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            z.extractall(tmp)

        settings = tmp / "word" / "settings.xml"
        settings.parent.mkdir(parents=True, exist_ok=True)

        parser = etree.XMLParser(remove_blank_text=False)
        if settings.exists():
            tree = etree.parse(str(settings), parser)
            root = tree.getroot()
        else:
            root = etree.Element(f"{{{W_NS}}}settings", nsmap={"w": W_NS})
            tree = etree.ElementTree(root)

        uf = root.find("w:updateFields", namespaces=NS)
        if uf is None:
            uf = etree.Element(f"{{{W_NS}}}updateFields")
            root.insert(0, uf)
        uf.set(f"{{{W_NS}}}val", "true")

        tree.write(str(settings), xml_declaration=True, encoding="UTF-8", standalone="yes")

        # Repack
        out = docx_path
        with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for p in tmp.rglob("*"):
                if p.is_dir():
                    continue
                z.write(p, p.relative_to(tmp).as_posix())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _clear_paragraph(p) -> None:
    for r in list(p.runs)[::-1]:
        p._p.remove(r._r)


def _add_toc_field(paragraph, levels: str = "1-3") -> None:
    """Insert a TOC field into an existing paragraph."""

    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    # Typical field switches: \o = levels, \h = hyperlinks, \z = hide page numbers in web layout, \u = use applied outline levels
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    # Placeholder display text before update
    t = OxmlElement("w:t")
    t.text = "(TOC will populate after updating fields)"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(t)
    r._r.append(fld_end)


def main() -> None:
    ap = argparse.ArgumentParser(description="Insert a Word TOC field at a placeholder paragraph.")
    ap.add_argument("docx", type=Path)
    ap.add_argument("--out", type=Path, default=None, help="Output DOCX path (default: in-place)")
    ap.add_argument("--placeholder", default="[[TOC]]", help="Paragraph text token to replace")
    ap.add_argument("--levels", default="1-3", help='Heading levels range, e.g. "1-3"')
    args = ap.parse_args()

    in_path = args.docx
    if not in_path.exists():
        raise FileNotFoundError(in_path)
    out_path = args.out or in_path
    if out_path != in_path:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(in_path, out_path)

    doc = Document(str(out_path))
    target = None
    for p in doc.paragraphs:
        if (p.text or "").strip() == args.placeholder:
            target = p
            break
    if target is None:
        raise RuntimeError(
            f"Could not find a paragraph whose text == placeholder {args.placeholder!r}. "
            "Tip: add a single paragraph containing that token."
        )

    _clear_paragraph(target)
    _add_toc_field(target, levels=args.levels)
    doc.save(str(out_path))

    _set_update_fields_on_open(out_path)
    print(f"[OK] Inserted TOC field at placeholder {args.placeholder!r} → {out_path}")
    print("Next: open in Word → Ctrl+A → F9 (Update Fields) → save → re-render.")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        if os.environ.get("DOCS_DEBUG") == "1":
            raise
        print(f"[ERROR] {e}", file=sys.stderr)
        raise SystemExit(2)
