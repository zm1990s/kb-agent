#!/usr/bin/env python3
"""Create reproducible DOCX fixtures for edge-case testing.

Why
---
Some DOCX edge cases (tracked changes, VML watermarks) are painful to reproduce
manually. This script generates small DOCX files and then injects minimal OOXML
so downstream skills can be tested deterministically.

This is intentionally *not* a general-purpose DOCX generator; it emits a tiny
set of fixtures with well-known markers.

Outputs
-------
- tracked_changes_fixture.docx
- watermark_fixture.docx

Usage
-----
python scripts/make_fixtures.py --outdir /mnt/data/fixtures
python scripts/make_fixtures.py --outdir /mnt/data/fixtures --only tracked

Notes
-----
- Tracked changes are represented by w:ins / w:del wrappers in document.xml.
- Watermarks are commonly VML in headers (w:pict + v:shape + v:textpath).
  LibreOffice headless may not render all VML variants; the watermark fixture is
  primarily for OOXML auditing/removal testing.
"""

from __future__ import annotations

import argparse
import os
import re
import zipfile
from datetime import datetime, timezone

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

NS = {"w": W_NS, "v": V_NS, "o": O_NS, "r": R_NS}


def _utc_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _read_zip_xml(z: zipfile.ZipFile, name: str) -> etree._Element:
    data = z.read(name)
    return etree.fromstring(data)


def _write_zip_xml(zout: zipfile.ZipFile, name: str, root: etree._Element) -> None:
    xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    zout.writestr(name, xml)


def _copy_zip_with_overrides(src_path: str, dst_path: str, overrides: dict[str, bytes]) -> None:
    """Copy all files from src zip to dst zip, overriding specific names."""
    with (
        zipfile.ZipFile(src_path, "r") as zin,
        zipfile.ZipFile(dst_path, "w", zipfile.ZIP_DEFLATED) as zout,
    ):
        for info in zin.infolist():
            name = info.filename
            if name in overrides:
                zout.writestr(name, overrides[name])
            else:
                zout.writestr(name, zin.read(name))
        # Add any new files not present in src
        for name, data in overrides.items():
            if name not in {i.filename for i in zin.infolist()}:
                zout.writestr(name, data)


def make_tracked_changes_fixture(out_path: str) -> None:
    """Create a minimal tracked-changes DOCX."""
    tmp_base = out_path + ".base.docx"

    doc = Document()
    doc.add_heading("Tracked Changes Fixture", level=1)
    p = doc.add_paragraph("Sentence with a pending edit: ")
    p.add_run("REPLACE_ME")
    doc.add_paragraph("End of fixture.")
    doc.save(tmp_base)

    with zipfile.ZipFile(tmp_base, "r") as z:
        doc_xml = _read_zip_xml(z, "word/document.xml")

    # Find the run containing 'REPLACE_ME'
    runs = doc_xml.xpath(".//w:r[w:t='REPLACE_ME']", namespaces=NS)
    if not runs:
        raise RuntimeError("Could not find marker text REPLACE_ME in generated document.xml")

    r = runs[0]
    parent = r.getparent()
    idx = parent.index(r)

    # Build <w:del> wrapper containing old text
    del_el = etree.Element(f"{{{W_NS}}}del", nsmap={"w": W_NS, "r": R_NS})
    del_el.set(f"{{{W_NS}}}id", "1")
    del_el.set(f"{{{W_NS}}}author", "FixtureMaker")
    del_el.set(f"{{{W_NS}}}date", _utc_now())
    del_r = etree.SubElement(del_el, f"{{{W_NS}}}r")
    del_t = etree.SubElement(del_r, f"{{{W_NS}}}delText")
    del_t.text = "REPLACE_ME"

    # Build <w:ins> wrapper containing new text
    ins_el = etree.Element(f"{{{W_NS}}}ins", nsmap={"w": W_NS, "r": R_NS})
    ins_el.set(f"{{{W_NS}}}id", "2")
    ins_el.set(f"{{{W_NS}}}author", "FixtureMaker")
    ins_el.set(f"{{{W_NS}}}date", _utc_now())
    ins_r = etree.SubElement(ins_el, f"{{{W_NS}}}r")
    ins_t = etree.SubElement(ins_r, f"{{{W_NS}}}t")
    ins_t.text = "INSERTED_TEXT"

    # Replace original run with del+ins
    parent.remove(r)
    parent.insert(idx, del_el)
    parent.insert(idx + 1, ins_el)

    # Enable trackRevisions in settings.xml
    with zipfile.ZipFile(tmp_base, "r") as z:
        settings = _read_zip_xml(z, "word/settings.xml")

    # Remove existing trackRevisions if present, then add
    for el in settings.xpath(".//w:trackRevisions", namespaces=NS):
        el.getparent().remove(el)
    settings.append(etree.Element(f"{{{W_NS}}}trackRevisions"))

    # Write overridden parts
    doc_xml_bytes = etree.tostring(
        doc_xml, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )
    settings_bytes = etree.tostring(
        settings, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    _copy_zip_with_overrides(
        tmp_base,
        out_path,
        {
            "word/document.xml": doc_xml_bytes,
            "word/settings.xml": settings_bytes,
        },
    )

    os.remove(tmp_base)


def _inject_vml_watermark_into_header(header_xml: etree._Element, text: str) -> None:
    """Inject a simple VML textpath watermark-like shape into the first header."""
    body = header_xml.find("w:body", namespaces=NS)
    if body is None:
        body = header_xml

    p = etree.Element(f"{{{W_NS}}}p")
    r = etree.SubElement(p, f"{{{W_NS}}}r")
    pict = etree.SubElement(r, f"{{{W_NS}}}pict")

    # <v:shape> with <v:textpath string="...">
    shape = etree.SubElement(pict, f"{{{V_NS}}}shape")
    shape.set("id", "PowerPlusWaterMarkObject1")
    shape.set(f"{{{O_NS}}}spid", "_x0000_s1025")
    shape.set("type", "#_x0000_t136")
    shape.set(
        "style",
        "position:absolute;margin-left:0;margin-top:0;width:468pt;height:234pt;rotation:315",
    )
    shape.set("fillcolor", "#d0d0d0")
    shape.set("stroked", "f")

    fill = etree.SubElement(shape, f"{{{V_NS}}}fill")
    fill.set("opacity", ".25")

    textpath = etree.SubElement(shape, f"{{{V_NS}}}textpath")
    textpath.set("style", "font-family:'Calibri';font-size:1pt")
    textpath.set("string", text)

    # Ensure the header has at least one paragraph; insert first.
    children = list(body)
    if children:
        body.insert(0, p)
    else:
        body.append(p)


def make_watermark_fixture(out_path: str) -> None:
    """Create a DOCX with a VML watermark-like shape in the header."""
    tmp_base = out_path + ".base.docx"

    doc = Document()
    # Force a header part to exist in the generated DOCX
    section = doc.sections[0]
    header = section.header
    if header.paragraphs:
        header.paragraphs[0].text = ""
    else:
        header.add_paragraph("")

    doc.add_heading("Watermark Fixture", level=1)
    doc.add_paragraph("This file contains a VML watermark-like header shape with text DRAFT.")
    doc.add_paragraph("LibreOffice headless may not render it; use OOXML audit/removal.")
    doc.save(tmp_base)

    with zipfile.ZipFile(tmp_base, "r") as z:
        # Find a header part to patch; prefer header1.xml
        header_name = None
        for name in z.namelist():
            if re.fullmatch(r"word/header\d+\.xml", name):
                header_name = name
                break
        if header_name is None:
            # Some docs might not have explicit headers; create header1.xml is non-trivial.
            raise RuntimeError("No header part found in generated DOCX; cannot inject watermark")
        header_xml = _read_zip_xml(z, header_name)

    _inject_vml_watermark_into_header(header_xml, text="DRAFT")
    header_bytes = etree.tostring(
        header_xml, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    _copy_zip_with_overrides(tmp_base, out_path, {header_name: header_bytes})
    os.remove(tmp_base)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--outdir", required=True, help="Directory to write fixtures")
    ap.add_argument(
        "--only",
        choices=["tracked", "watermark", "all"],
        default="all",
        help="Which fixtures to generate",
    )
    args = ap.parse_args()

    os.makedirs(args.outdir, exist_ok=True)

    if args.only in ("tracked", "all"):
        out = os.path.join(args.outdir, "tracked_changes_fixture.docx")
        make_tracked_changes_fixture(out)
        print(f"[OK] wrote {out}")

    if args.only in ("watermark", "all"):
        out = os.path.join(args.outdir, "watermark_fixture.docx")
        make_watermark_fixture(out)
        print(f"[OK] wrote {out}")


if __name__ == "__main__":
    main()
