#!/usr/bin/env python3
"""DOCX patch helper for features missing in python-docx.

Supports:
- enabling Track Revisions flag
- producing true tracked changes (convert an existing <w:ins> into <w:del> + insert new <w:ins>)
- adding true Word comments (comments.xml + anchors + rels + content types)
- adding a hyperlink / header date / page number field (via python-docx low-level XML)

This is intentionally pragmatic: it targets common automation tasks, not full OOXML generality.
"""

from __future__ import annotations

import argparse
import datetime as _dt
import os
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

from lxml import etree

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as docx_qn
except Exception:
    Document = None  # type: ignore

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

NS = {
    "w": W_NS,
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def _qn(prefix: str, local: str) -> str:
    # Build a Clark-notation qualified name for common OOXML namespaces.
    if prefix == "w":
        return f"{{{W_NS}}}{local}"
    if prefix == "pr":
        return f"{{{PKG_REL_NS}}}{local}"
    if prefix == "ct":
        return f"{{{CT_NS}}}{local}"
    if prefix == "r":
        return f"{{{NS['r']}}}{local}"
    raise ValueError(prefix)


def iso_now() -> str:
    # Word likes a Z suffix.
    return _dt.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def unzip_docx(docx_path: Path, out_dir: Path) -> None:
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(docx_path, "r") as z:
        z.extractall(out_dir)


def zip_docx(in_dir: Path, out_docx_path: Path) -> None:
    if out_docx_path.exists():
        out_docx_path.unlink()
    with zipfile.ZipFile(out_docx_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for root, _dirs, files in os.walk(in_dir):
            for f in files:
                abs_path = Path(root) / f
                rel_path = abs_path.relative_to(in_dir)
                z.write(abs_path, rel_path.as_posix())


def enable_track_revisions(unzipped: Path) -> None:
    settings_path = unzipped / "word" / "settings.xml"
    settings_path.parent.mkdir(parents=True, exist_ok=True)

    parser = etree.XMLParser(remove_blank_text=False)

    # Some minimal DOCX files omit settings.xml. Create a minimal settings part.
    if not settings_path.exists():
        root = etree.Element(_qn("w", "settings"), nsmap={"w": W_NS})
        etree.ElementTree(root).write(
            str(settings_path), xml_declaration=True, encoding="UTF-8", standalone="yes"
        )

    tree = etree.parse(str(settings_path), parser)
    root = tree.getroot()
    track = root.find("w:trackRevisions", namespaces=NS)
    if track is None:
        track = etree.Element(_qn("w", "trackRevisions"))
        compat = root.find("w:compat", namespaces=NS)
        if compat is not None:
            compat.addnext(track)
        else:
            root.insert(0, track)
    tree.write(str(settings_path), xml_declaration=True, encoding="UTF-8", standalone="yes")


def tracked_replace_ins_id(
    unzipped: Path,
    ins_w_id: str,
    new_text: str,
    del_w_id: str = "auto",
    new_ins_w_id: str = "auto",
    author: str = "ChatGPT",
    date_iso: str | None = None,
) -> None:
    """Convert an existing <w:ins w:id=ins_w_id> into <w:del> and insert a new <w:ins> after it."""

    date_iso = date_iso or iso_now()
    doc_xml = unzipped / "word" / "document.xml"
    parser = etree.XMLParser(remove_blank_text=False)
    tree = etree.parse(str(doc_xml), parser)
    root = tree.getroot()

    ins = root.find(f".//w:ins[@w:id='{ins_w_id}']", namespaces=NS)
    if ins is None:
        raise RuntimeError(f"Could not find <w:ins w:id='{ins_w_id}'> in document.xml")

    # Pick IDs that won't collide with existing w:id values.
    existing_ids: set[int] = set()
    for el in root.iter():
        # Only consider the WordprocessingML w:id attribute.
        v = el.get(_qn("w", "id"))
        if v is None:
            continue
        try:
            existing_ids.add(int(v))
        except Exception:
            continue
    max_id = max(existing_ids) if existing_ids else 0

    if del_w_id == "auto":
        del_w_id = str(max_id + 1)
    if new_ins_w_id == "auto":
        # Ensure different from deletion id even if user set del_w_id explicitly.
        try:
            base = max(max_id, int(del_w_id))
        except Exception:
            base = max_id
        new_ins_w_id = str(base + 1)

    # Turn ins into del.
    ins.tag = _qn("w", "del")
    ins.attrib[_qn("w", "id")] = del_w_id
    ins.attrib[_qn("w", "author")] = author
    ins.attrib[_qn("w", "date")] = date_iso

    # Rename all w:t to w:delText
    for t in ins.findall(".//w:t", namespaces=NS):
        t.tag = _qn("w", "delText")

    parent = ins.getparent()
    pos = parent.index(ins)

    new_ins = etree.Element(_qn("w", "ins"))
    new_ins.attrib[_qn("w", "id")] = new_ins_w_id
    new_ins.attrib[_qn("w", "author")] = author
    new_ins.attrib[_qn("w", "date")] = date_iso

    r = etree.SubElement(new_ins, _qn("w", "r"))
    t = etree.SubElement(r, _qn("w", "t"))
    t.text = new_text

    parent.insert(pos + 1, new_ins)
    tree.write(str(doc_xml), xml_declaration=True, encoding="UTF-8", standalone="yes")


def add_comment_to_first_indented_paragraph(
    unzipped: Path,
    comment_text: str,
    indent_left_twips: int | None = None,
    contains: str | None = None,
    comment_id: str = "auto",
    author: str = "ChatGPT",
    date_iso: str | None = None,
) -> None:
    """Add a Word comment anchored to a paragraph.

    Anchor selection:
      - If neither `indent_left_twips` nor `contains` is provided, anchors to the **first non-empty**
        paragraph (fallback: first paragraph).
      - If `indent_left_twips` is provided, the target paragraph must have `w:pPr/w:ind/@w:left` equal
        to that value (twips). (Note: many paragraphs don't have explicit indentation.)
      - If `contains` is provided, the target paragraph's visible text must contain that substring.

    `comment_id`:
      - "auto" (default) selects the next non-colliding id across existing anchors and comments.xml.

    This function edits the unzipped DOCX in place (document.xml, comments.xml, rels, content types).
    """

    date_iso = date_iso or iso_now()
    parser = etree.XMLParser(remove_blank_text=False)

    doc_xml = unzipped / "word" / "document.xml"
    tree = etree.parse(str(doc_xml), parser)
    root = tree.getroot()

    paras = list(root.findall(".//w:p", namespaces=NS))
    if not paras:
        raise RuntimeError("No paragraphs found in document.xml; cannot anchor comment.")

    def para_text(p: etree._Element) -> str:
        return "".join([(x.text or "") for x in p.findall(".//w:t", namespaces=NS)])

    target_p = None

    if indent_left_twips is None and contains is None:
        for p_el in paras:
            if para_text(p_el).strip():
                target_p = p_el
                break
        if target_p is None:
            target_p = paras[0]
    else:
        want_indent = None if indent_left_twips is None else str(indent_left_twips)
        for p_el in paras:
            if want_indent is not None:
                pPr = p_el.find("w:pPr", namespaces=NS)
                if pPr is None:
                    continue
                ind = pPr.find("w:ind", namespaces=NS)
                if ind is None:
                    continue
                if ind.get(_qn("w", "left")) != want_indent:
                    continue
            if contains is not None:
                if contains not in para_text(p_el):
                    continue
            target_p = p_el
            break

        if target_p is None:
            bits = []
            if indent_left_twips is not None:
                bits.append(f"indent_left_twips={indent_left_twips}")
            if contains is not None:
                bits.append(f"contains={contains!r}")
            hint = (
                "Tip: omit both flags to anchor to the first paragraph, "
                "or use --contains to match a short snippet from the target paragraph."
            )
            raise RuntimeError(
                "Could not find paragraph matching predicate (" + ", ".join(bits) + ").\n" + hint
            )

    # Auto-pick a comment id if requested.
    if comment_id == "auto":
        used: set[int] = set()
        # Existing anchors in document.xml
        for el in root.findall(".//w:commentRangeStart", namespaces=NS) + root.findall(
            ".//w:commentRangeEnd", namespaces=NS
        ):
            v = el.get(_qn("w", "id"))
            if v is None:
                continue
            try:
                used.add(int(v))
            except Exception:
                continue
        # Existing comments.xml ids
        comments_path = unzipped / "word" / "comments.xml"
        if comments_path.exists():
            try:
                ctree = etree.parse(str(comments_path), parser)
                for c in ctree.getroot().findall("w:comment", namespaces=NS):
                    v = c.get(_qn("w", "id"))
                    if v is None:
                        continue
                    try:
                        used.add(int(v))
                    except Exception:
                        continue
            except Exception:
                pass
        new_id = 0
        while new_id in used:
            new_id += 1
        comment_id = str(new_id)

    # Ensure word/comments.xml exists.
    comments_path = unzipped / "word" / "comments.xml"
    if not comments_path.exists():
        comments_root = etree.Element(_qn("w", "comments"), nsmap={"w": W_NS})
        etree.ElementTree(comments_root).write(
            str(comments_path), xml_declaration=True, encoding="UTF-8", standalone="yes"
        )

    # Ensure relationship from document.xml -> comments.xml exists.
    rels_path = unzipped / "word" / "_rels" / "document.xml.rels"
    rels_path.parent.mkdir(parents=True, exist_ok=True)
    if rels_path.exists():
        rtree = etree.parse(str(rels_path), parser)
        rroot = rtree.getroot()
    else:
        rroot = etree.Element(_qn("pr", "Relationships"), nsmap={"pr": PKG_REL_NS})
        rtree = etree.ElementTree(rroot)

    # Find or create comments relationship.
    comments_rId = None
    for rel in rroot.findall(".//pr:Relationship", namespaces={"pr": PKG_REL_NS}):
        if (
            rel.get("Type")
            == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        ):
            comments_rId = rel.get("Id")
            break

    if comments_rId is None:
        # Pick next rId.
        max_rid = 0
        for rel in rroot.findall(".//pr:Relationship", namespaces={"pr": PKG_REL_NS}):
            rid = rel.get("Id", "")
            if rid.startswith("rId"):
                try:
                    max_rid = max(max_rid, int(rid[3:]))
                except Exception:
                    pass
        comments_rId = f"rId{max_rid + 1}"
        new_rel = etree.SubElement(rroot, _qn("pr", "Relationship"))
        new_rel.set("Id", comments_rId)
        new_rel.set(
            "Type",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
        )
        new_rel.set("Target", "comments.xml")
        rtree.write(str(rels_path), xml_declaration=True, encoding="UTF-8", standalone="yes")

    # Ensure content types for comments.xml
    ct_path = unzipped / "[Content_Types].xml"
    cttree = etree.parse(str(ct_path), parser)
    ctroot = cttree.getroot()
    have_override = False
    for ov in ctroot.findall(".//ct:Override", namespaces={"ct": CT_NS}):
        if ov.get("PartName") == "/word/comments.xml":
            have_override = True
            break
    if not have_override:
        ov = etree.SubElement(ctroot, _qn("ct", "Override"))
        ov.set("PartName", "/word/comments.xml")
        ov.set(
            "ContentType",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        )
        cttree.write(str(ct_path), xml_declaration=True, encoding="UTF-8", standalone="yes")

    # Add anchor elements around the paragraph's first run.
    runs = target_p.findall("w:r", namespaces=NS)
    if not runs:
        runs = [etree.SubElement(target_p, _qn("w", "r"))]

    crs = etree.Element(_qn("w", "commentRangeStart"))
    crs.set(_qn("w", "id"), comment_id)
    target_p.insert(target_p.index(runs[0]), crs)

    cre = etree.Element(_qn("w", "commentRangeEnd"))
    cre.set(_qn("w", "id"), comment_id)
    target_p.insert(target_p.index(runs[-1]) + 1, cre)

    cref_run = etree.Element(_qn("w", "r"))
    cref = etree.SubElement(cref_run, _qn("w", "commentReference"))
    cref.set(_qn("w", "id"), comment_id)
    target_p.insert(target_p.index(cre) + 1, cref_run)

    tree.write(str(doc_xml), xml_declaration=True, encoding="UTF-8", standalone="yes")

    # Add comment entry to comments.xml.
    ctree = etree.parse(str(comments_path), parser)
    croot = ctree.getroot()

    # Avoid duplicate ids.
    if croot.find(f".//w:comment[@w:id='{comment_id}']", namespaces=NS) is not None:
        raise RuntimeError(f"comments.xml already contains comment id={comment_id}")

    c = etree.SubElement(croot, _qn("w", "comment"))
    c.set(_qn("w", "id"), comment_id)
    c.set(_qn("w", "author"), author)
    c.set(_qn("w", "date"), date_iso)

    # <w:p><w:r><w:t>...</w:t></w:r></w:p>
    p_el = etree.SubElement(c, _qn("w", "p"))
    r_el = etree.SubElement(p_el, _qn("w", "r"))
    t_el = etree.SubElement(r_el, _qn("w", "t"))
    t_el.text = comment_text

    ctree.write(str(comments_path), xml_declaration=True, encoding="UTF-8", standalone="yes")


def require_python_docx() -> None:
    if Document is None:
        raise RuntimeError("python-docx is not available in this environment")


def add_header_date(docx_path: Path, date_text: str) -> None:
    require_python_docx()
    doc = Document(str(docx_path))
    sec = doc.sections[0]
    header = sec.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.text = date_text
    doc.save(str(docx_path))


def add_page_number_field_to_footer(docx_path: Path) -> None:
    require_python_docx()
    doc = Document(str(docx_path))
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(docx_qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(docx_qn("xml:space"), "preserve")
    instr.text = " PAGE \\* MERGEFORMAT "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(docx_qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(docx_qn("w:fldCharType"), "end")

    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(txt)
    r._r.append(fld_end)

    doc.save(str(docx_path))


def hyperlink_first_paragraph(docx_path: Path, url: str, underline: bool = True) -> None:
    require_python_docx()
    doc = Document(str(docx_path))
    if not doc.paragraphs:
        raise RuntimeError("Document has no paragraphs")

    p = doc.paragraphs[0]
    text = p.text or ""
    # Remove existing runs
    for run in list(p.runs)[::-1]:
        p._p.remove(run._r)

    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(docx_qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(docx_qn("w:val"), "0000FF")
    rPr.append(color)
    if underline:
        u = OxmlElement("w:u")
        u.set(docx_qn("w:val"), "single")
        rPr.append(u)
    r.append(rPr)

    t = OxmlElement("w:t")
    # Preserve leading/trailing spaces if present.
    if text[:1].isspace() or text[-1:].isspace() or "  " in text:
        t.set(docx_qn("xml:space"), "preserve")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    p._p.append(hyperlink)

    doc.save(str(docx_path))


# ---------------- CLI ----------------


def _cmd_ooxml(args: argparse.Namespace) -> None:
    # IMPORTANT:
    # If python-docx helpers ran with --out, they modified the *output* file, not the input.
    # The OOXML patch step must therefore unzip/patch the current working doc (out or in-place),
    # not always the original input path.
    docx_current = Path(args.out or args.docx)
    if not docx_current.exists():
        raise FileNotFoundError(docx_current)

    out = docx_current
    workdir = Path(tempfile.mkdtemp(prefix="docx_patch_"))
    unz = workdir / "unz"
    unzip_docx(docx_current, unz)

    if args.enable_track:
        enable_track_revisions(unz)

    if args.tracked_replace_ins_id is not None:
        tracked_replace_ins_id(
            unzipped=unz,
            ins_w_id=args.tracked_replace_ins_id,
            new_text=args.new_text,
            del_w_id=args.del_id,
            new_ins_w_id=args.ins_id,
            author=args.author,
            date_iso=args.date,
        )

    if args.add_comment:
        add_comment_to_first_indented_paragraph(
            unzipped=unz,
            comment_text=args.comment_text,
            indent_left_twips=args.indent_left_twips,
            contains=args.contains,
            comment_id=args.comment_id,
            author=args.author,
            date_iso=args.date,
        )

    zip_docx(unz, out)
    shutil.rmtree(workdir, ignore_errors=True)


def _print_ok_summary(out_path: Path, ops: list[str]) -> None:
    print(f"[OK] Patched -> {out_path}")
    for op in ops:
        print(f"  - {op}")


def main() -> int:
    ap = argparse.ArgumentParser(description="DOCX patch helper (OOXML + python-docx utilities)")
    ap.add_argument("docx", help="Input DOCX")
    ap.add_argument("--out", default=None, help="Output DOCX (default: overwrite input)")

    # OOXML switches
    ap.add_argument(
        "--enable-track",
        action="store_true",
        help="Enable <w:trackRevisions/> in settings.xml",
    )
    ap.add_argument("--author", default="ChatGPT", help="Author metadata")
    ap.add_argument("--date", default=None, help="ISO date for metadata (default: now)")

    ap.add_argument(
        "--tracked-replace-ins-id",
        default=None,
        help="Convert <w:ins w:id=...> into delete+insert",
    )
    ap.add_argument("--new-text", default="", help="New text for tracked insertion")
    ap.add_argument(
        "--del-id",
        default="auto",
        help="w:id for the deletion (default: auto-pick a non-colliding id)",
    )
    ap.add_argument(
        "--ins-id",
        default="auto",
        help="w:id for the new insertion (default: auto-pick a non-colliding id)",
    )

    ap.add_argument(
        "--add-comment", action="store_true", help="Add a Word comment via comments.xml"
    )
    ap.add_argument("--comment-text", default="", help="Comment body")
    ap.add_argument(
        "--indent-left-twips",
        type=int,
        default=None,
        help="Match paragraphs by indentation left (twips)",
    )
    ap.add_argument("--contains", default=None, help="Require substring in paragraph text")
    ap.add_argument(
        "--comment-id",
        default="auto",
        help="Comment id (default: auto-pick a non-colliding id)",
    )

    # python-docx helpers
    ap.add_argument(
        "--header-date",
        default=None,
        help="Set first section header to this text (right aligned)",
    )
    ap.add_argument(
        "--add-page-numbers",
        action="store_true",
        help="Add a centered page number field to footer",
    )
    ap.add_argument(
        "--hyperlink-first",
        default=None,
        help="Add hyperlink to first paragraph text to this URL",
    )

    args = ap.parse_args()

    ops: list[str] = []

    # First apply python-docx changes (they rewrite OOXML parts), then patch OOXML.
    docx_path = Path(args.out or args.docx)
    if args.out:
        shutil.copyfile(args.docx, args.out)

    if args.header_date is not None or args.add_page_numbers or args.hyperlink_first is not None:
        require_python_docx()
        if args.header_date is not None:
            add_header_date(docx_path, args.header_date)
            ops.append(f"header-date={args.header_date!r}")
        if args.add_page_numbers:
            add_page_number_field_to_footer(docx_path)
            ops.append("add-page-numbers")
        if args.hyperlink_first is not None:
            hyperlink_first_paragraph(docx_path, args.hyperlink_first)
            ops.append(f"hyperlink-first={args.hyperlink_first!r}")

    # OOXML patch pass
    if args.enable_track or args.tracked_replace_ins_id or args.add_comment:
        if args.enable_track:
            ops.append("enable-track")
        if args.tracked_replace_ins_id:
            ops.append(f"tracked-replace-ins-id={args.tracked_replace_ins_id!r}")
        if args.add_comment:
            ops.append("add-comment")
        _cmd_ooxml(args)

    if ops:
        _print_ok_summary(Path(args.out or args.docx), ops)
    else:
        print("[OK] No changes requested (nothing to do)")

    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as e:
        if os.environ.get("DOCS_DEBUG") == "1":
            raise
        print(f"[ERROR] {e}", file=sys.stderr)
        raise SystemExit(2)
