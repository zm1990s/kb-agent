#!/usr/bin/env python3
"""Convert an XLSX worksheet to a simple DOCX table.

Scope
-----
This is a "minimum useful" converter intended for reports:
- Reads values from a sheet
- Creates a DOCX with a grid table
- Applies a few common-sense formats (header bold, basic alignment)
- Uses heuristic column widths based on max content length

Limits
------
- It does NOT reproduce Excel formulas, merged cells, conditional formatting, charts, or complex number formats.
- Number formatting is best-effort. If you need exact appearance, consider exporting the range as an image instead.
"""

from __future__ import annotations

import argparse
import datetime as _dt
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from openpyxl import load_workbook


def _mark_row_as_header(row) -> None:
    """Mark a table row as a repeating header row (w:tblHeader).

    This improves accessibility and ensures header rows repeat when the table
    spans multiple pages.
    """

    tr = row._tr  # python-docx internal
    trPr = tr.get_or_add_trPr()
    # Avoid duplicating if re-run
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def _is_empty(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str) and not v.strip():
        return True
    return False


def _best_effort_format(cell) -> str:
    v = cell.value
    if v is None:
        return ""
    if isinstance(v, (bool, int)):
        return str(v)
    if isinstance(v, float):
        fmt = (cell.number_format or "").lower()
        if "%" in fmt:
            return f"{v * 100:.2f}%"
        if "$" in fmt:
            return f"${v:,.2f}"
        return f"{v:g}"
    if isinstance(v, (_dt.datetime, _dt.date)):
        try:
            return v.isoformat()
        except Exception:
            return str(v)
    return str(v)


def _used_bounds(ws):
    # Compute a tight-ish used range; ws.max_row/col may include trailing empties.
    max_r = 0
    max_c = 0
    for row in ws.iter_rows():
        for cell in row:
            if not _is_empty(cell.value):
                max_r = max(max_r, cell.row)
                max_c = max(max_c, cell.column)
    return max_r, max_c


def main() -> None:
    ap = argparse.ArgumentParser(description="Convert an XLSX sheet to a DOCX table")
    ap.add_argument("xlsx", type=Path)
    # Keep the original positional output path for backwards compatibility, but also
    # support an --out alias (agents often assume this flag exists).
    ap.add_argument("out_docx", type=Path, nargs="?", default=None)
    ap.add_argument("--out", type=Path, default=None, help="Output DOCX path (alias)")
    ap.add_argument("--sheet", default=None, help="Sheet name (default: active)")
    ap.add_argument("--header_rows", type=int, default=1, help="How many top rows to bold")
    ap.add_argument("--title", default=None, help="Optional title paragraph above the table")
    args = ap.parse_args()

    out_path = args.out or args.out_docx
    if out_path is None:
        raise RuntimeError("Missing output path. Provide a positional out_docx or --out.")

    if not args.xlsx.exists():
        raise FileNotFoundError(args.xlsx)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    wb = load_workbook(args.xlsx, data_only=True)
    ws = wb[args.sheet] if args.sheet else wb.active

    max_r, max_c = _used_bounds(ws)
    if max_r == 0 or max_c == 0:
        raise RuntimeError("Sheet appears empty")

    # Gather formatted strings for width heuristics
    data = []
    col_max_len = [0] * max_c
    for r in range(1, max_r + 1):
        row_vals = []
        for c in range(1, max_c + 1):
            cell = ws.cell(row=r, column=c)
            s = _best_effort_format(cell)
            row_vals.append(s)
            col_max_len[c - 1] = max(col_max_len[c - 1], len(s))
        data.append(row_vals)

    doc = Document()
    if args.title:
        doc.add_paragraph(args.title)

    table = doc.add_table(rows=max_r, cols=max_c)
    table.style = "Table Grid"
    table.autofit = False

    # Column widths (heuristic): ~0.12in per char, clamped
    col_widths = []
    for ml in col_max_len:
        w = max(0.8, min(3.0, 0.12 * max(ml, 4)))
        col_widths.append(w)

    for r in range(max_r):
        for c in range(max_c):
            cell = ws.cell(row=r + 1, column=c + 1)
            out = table.cell(r, c)
            out.width = Inches(col_widths[c])
            p = out.paragraphs[0]
            p.text = data[r][c]

            # Alignment best-effort
            horiz = getattr(cell.alignment, "horizontal", None)
            if horiz in {"right", "center", "left"}:
                p.alignment = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                }[horiz]

            # Bold header rows
            if r < args.header_rows and p.runs:
                for run in p.runs:
                    run.bold = True

        # Mark header rows in OOXML for a11y + repeating headers
        if r < args.header_rows:
            _mark_row_as_header(table.rows[r])

    doc.save(str(out_path))
    print(f"[OK] Wrote {out_path} (rows={max_r}, cols={max_c}, sheet={ws.title!r})")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        if os.environ.get("DOCS_DEBUG") == "1":
            raise
        print(f"[ERROR] {e}", file=sys.stderr)
        raise SystemExit(2)
