#!/usr/bin/env python3
"""Add a simple VML watermark-like object into a document header.

Why this exists
--------------
The v6 toolkit includes `watermark_audit_remove.py` (removal) but stress-testing
watermark workflows needs a way to *add* a watermark as well.

Implementation notes
-------------------
- Word watermarks are commonly VML shapes embedded in headers. This script adds a
  conservative VML `<v:shape>` containing a `<v:textpath string="..."/>`.
- If the input docx has no header part yet, we create a minimal header using
  python-docx (empty header text) and then patch that header.
- Rendering differs between Word/LibreOffice; the goal is deterministic OOXML
  that is discoverable/removable by `watermark_audit_remove.py`.

Usage
-----
python scripts/watermark_add.py in.docx --out out.docx --text "CONFIDENTIAL"
"""

from __future__ import annotations

import argparse
import re
import tempfile
import zipfile
from pathlib import Path

from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"

NS = {"w": W_NS, "v": V_NS, "o": O_NS}


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def v(tag: str) -> str:
    return f"{{{V_NS}}}{tag}"


def o(tag: str) -> str:
    return f"{{{O_NS}}}{tag}"


def _xml_bytes(root: etree._Element) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _pick_header_part(names: list[str]) -> str | None:
    headers = [n for n in names if re.fullmatch(r"word/header\d+\.xml", n)]
    if not headers:
        return None
    headers_sorted = sorted(headers, key=lambda s: int(re.findall(r"\d+", s)[0]))
    return headers_sorted[0]


def _make_watermark_paragraph(text: str) -> etree._Element:
    # A minimal VML shape. Word-specific style fields are best-effort.
    p = etree.Element(w("p"))
    r_el = etree.SubElement(p, w("r"))
    pict = etree.SubElement(r_el, w("pict"))

    shape = etree.SubElement(
        pict,
        v("shape"),
        {
            "id": "DocxSkillWatermark",
            o("spid"): "_x0000_s1025",
            "type": "#_x0000_t136",
            "style": (
                "position:absolute;"
                "margin-left:0;margin-top:0;"
                "width:468pt;height:468pt;"
                "rotation:315;"
                "z-index:-251654144;"
                "mso-position-horizontal:center;"
                "mso-position-vertical:center;"
                "mso-wrap-edited:f;"
            ),
            "fillcolor": "#C0C0C0",
            "stroked": "f",
        },
    )

    # Semi-transparent fill.
    etree.SubElement(shape, v("fill"), {"opacity": "0.15"})

    # Text path.
    etree.SubElement(
        shape,
        v("textpath"),
        {"style": 'font-family:"Calibri";font-size:1pt', "string": text},
    )
    etree.SubElement(shape, v("path"), {"textpathok": "t"})
    return p


def _ensure_header_exists(in_docx: str) -> str:
    """Return a path to a docx that is guaranteed to have a header part.

    If the input already has a header part, returns the original path.
    Otherwise, writes a minimal-header variant to a temp file and returns that.
    """
    with zipfile.ZipFile(in_docx, "r") as z:
        if _pick_header_part(z.namelist()) is not None:
            return in_docx

    from docx import Document  # local import to keep deps light for simple uses

    doc = Document(in_docx)
    sec = doc.sections[0]
    hdr = sec.header
    if not hdr.paragraphs:
        hdr.add_paragraph("")
    # Add a single whitespace run so the header part is materialized, while
    # remaining visually blank.
    if not (hdr.paragraphs[0].text or "").strip():
        hdr.paragraphs[0].add_run(" ")

    td = tempfile.mkdtemp(prefix="docx_wm_header_")
    out = str(Path(td) / "with_header.docx")
    doc.save(out)
    return out


def add_watermark(in_docx: str, out_docx: str, text: str) -> None:
    patched_input = _ensure_header_exists(in_docx)
    with zipfile.ZipFile(patched_input, "r") as zin:
        names = zin.namelist()
        header = _pick_header_part(names)
        if header is None:
            raise SystemExit("[watermark_add] Could not create or locate a header part")

        overrides: dict[str, bytes] = {}
        root = etree.fromstring(zin.read(header))

        # Append our watermark paragraph at the end of the header.
        root.append(_make_watermark_paragraph(text))
        overrides[header] = _xml_bytes(root)

        with zipfile.ZipFile(out_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                name = info.filename
                if name in overrides:
                    zout.writestr(name, overrides[name])
                else:
                    zout.writestr(name, zin.read(name))

    print(f"[OK] wrote {out_docx} (patched {header})")


def main() -> None:
    ap = argparse.ArgumentParser(description="Add a VML watermark-like object to a DOCX header")
    ap.add_argument("in_docx")
    ap.add_argument("--out", required=True)
    ap.add_argument("--text", required=True, help="Watermark string")
    args = ap.parse_args()

    add_watermark(args.in_docx, args.out, args.text)


if __name__ == "__main__":
    main()
