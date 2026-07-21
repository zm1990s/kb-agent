#!/usr/bin/env python3
"""Normalize styles / reduce formatting drift.

This is a conservative companion to style_lint.py.

Default behavior (safe-ish):
- Clear *run-level* direct formatting overrides (bold/italic/underline,
  font name/size/color) by setting them to None, letting paragraph/style
  definitions drive appearance.

Optional behaviors:
- Clear paragraph formatting overrides (space before/after, indents)
- Enforce simple heading spacing rules

Always render to PNGs after normalization, because clearing overrides can change layout.

CLI compatibility
-----------------
Historically this script used positional args:

  python scripts/style_normalize.py in.docx out.docx

During stress-testing it's easy to assume a common `--out` flag exists. v8 keeps
the original signature but also supports `--out` as an alias:

  python scripts/style_normalize.py in.docx --out out.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def _is_heading(style_name: str | None) -> bool:
    if not style_name:
        return False
    return style_name.lower().startswith("heading")


def clear_run_direct_formatting(doc: Document) -> int:
    changed = 0
    for p in doc.paragraphs:
        for r in p.runs:
            f = r.font
            # Any non-None here is a direct override.
            if (
                f.name is not None
                or f.size is not None
                or f.bold is not None
                or f.italic is not None
                or f.underline is not None
                or f.color.rgb is not None
            ):
                f.name = None
                f.size = None
                f.bold = None
                f.italic = None
                f.underline = None
                f.color.rgb = None
                changed += 1

    # runs in tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        f = r.font
                        if (
                            f.name is not None
                            or f.size is not None
                            or f.bold is not None
                            or f.italic is not None
                            or f.underline is not None
                            or f.color.rgb is not None
                        ):
                            f.name = None
                            f.size = None
                            f.bold = None
                            f.italic = None
                            f.underline = None
                            f.color.rgb = None
                            changed += 1
    return changed


def clear_paragraph_direct_formatting(doc: Document) -> int:
    changed = 0

    def _clear(p):
        nonlocal changed
        pf = p.paragraph_format
        # Indents
        for attr in (
            "left_indent",
            "right_indent",
            "first_line_indent",
        ):
            if getattr(pf, attr) is not None:
                setattr(pf, attr, None)
                changed += 1
        # Spacing
        for attr in (
            "space_before",
            "space_after",
            "line_spacing",
            "line_spacing_rule",
        ):
            if getattr(pf, attr) is not None:
                setattr(pf, attr, None)
                changed += 1

    for p in doc.paragraphs:
        _clear(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _clear(p)
    return changed


def enforce_heading_spacing(doc: Document, space_after_pt: float = 6.0) -> int:
    from docx.shared import Pt

    changed = 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style is not None else None
        if _is_heading(sname):
            if p.paragraph_format.space_after != Pt(space_after_pt):
                p.paragraph_format.space_after = Pt(space_after_pt)
                changed += 1
    return changed


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("input_docx")
    # Backwards compatible positional output.
    ap.add_argument("output_docx", nargs="?", help="Output DOCX path (positional)")
    # Alias used across the rest of this skill.
    ap.add_argument(
        "--out",
        default=None,
        help="Output DOCX path (alias for positional output_docx)",
    )
    ap.add_argument(
        "--clear_paragraph_format",
        action="store_true",
        help="Also clear paragraph-level direct formatting overrides (more invasive).",
    )
    ap.add_argument(
        "--enforce_heading_spacing",
        action="store_true",
        help="Enforce a simple heading spacing rule (space-after = 6pt).",
    )
    args = ap.parse_args()

    out_arg = args.out or args.output_docx
    if not out_arg:
        ap.error("output_docx (positional) or --out is required")

    doc = Document(args.input_docx)
    run_changes = clear_run_direct_formatting(doc)
    para_changes = 0
    heading_changes = 0

    if args.clear_paragraph_format:
        para_changes = clear_paragraph_direct_formatting(doc)
    if args.enforce_heading_spacing:
        heading_changes = enforce_heading_spacing(doc)

    out = Path(out_arg)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    print(
        f"[OK] wrote {out} | run_overrides_cleared={run_changes} "
        f"para_overrides_cleared={para_changes} heading_spacing_updates={heading_changes}"
    )


if __name__ == "__main__":
    main()
