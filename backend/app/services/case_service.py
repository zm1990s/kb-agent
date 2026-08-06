"""Case 录入导出服务：Tiptap JSON → .docx / .pdf + 纯文本。

Tiptap 文档结构（简化）：
  {type: "doc", content: [node, ...]}
  块级 type ∈ heading|paragraph|bulletList|orderedList|blockquote|codeBlock|
               horizontalRule|image|table
  内联：node.content = [
    {type:"text", text, marks:[{type:"bold"|"italic"|"strike"|"code"|"link", attrs}]},
    {type:"hardBreak"},
    ...
  ]
  图片：{type:"image", attrs:{src: "data:image/png;base64,..."}}

图片按 base64 dataURL 内嵌（首版决策）；非 data: 的图片 src 忽略（不外链拉取）。
"""

import base64
import io
import logging

logger = logging.getLogger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

# 单图解码上限（防超大 base64 拖垮导出）
_MAX_IMAGE_BYTES = 8 * 1024 * 1024


# ── 解析 base64 图片 ────────────────────────────────────────────────────────
def _decode_data_url(src: str) -> bytes | None:
    """解析 data:image/*;base64,xxx，返回字节；非 data URL 或超限返回 None。"""
    if not src.startswith("data:"):
        return None
    try:
        header, b64 = src.split(",", 1)
        if "base64" not in header:
            return None
        data = base64.b64decode(b64)
    except (ValueError, base64.binascii.Error):  # type: ignore[attr-defined]
        return None
    if len(data) > _MAX_IMAGE_BYTES:
        logger.warning("case 导出跳过超大内嵌图片 %d bytes", len(data))
        return None
    return data


# ── 纯文本提取（供 content_text 全文检索） ──────────────────────────────────
def _node_text(node: dict) -> str:
    if node.get("type") == "text":
        return node.get("text", "")
    parts = [_node_text(c) for c in node.get("content", []) or []]
    return "".join(parts)


def _extract_plain_text(doc: dict) -> str:
    lines: list[str] = []
    for node in doc.get("content", []) or []:
        t = node.get("type")
        if t in ("bulletList", "orderedList"):
            for item in node.get("content", []) or []:
                lines.append(_node_text(item))
        elif t == "image":
            continue
        else:
            txt = _node_text(node)
            if txt:
                lines.append(txt)
    return "\n".join(lines).strip()


# ── DOCX 生成 ───────────────────────────────────────────────────────────────
_DOCX_FONT = "Microsoft YaHei"


def _set_doc_font(doc: object, font_name: str = _DOCX_FONT) -> None:
    """设置 docx 文档默认字体（ASCII + 东亚字体槽），使中文正确渲染。"""
    from docx.oxml.ns import qn

    rPr = doc.styles["Normal"].element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """在 docx 段落中插入真实超链接。"""
    import docx.opc.constants as opc_const
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(url, opc_const.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    new_run.append(t_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _run_inline(paragraph, node: dict) -> None:
    """把内联节点写入 docx 段落。
    支持 bold / italic / strike / code marks、link mark（真实超链接）、hardBreak。
    """
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    for child in node.get("content", []) or []:
        ct = child.get("type")
        if ct == "hardBreak":
            br_run = paragraph.add_run()
            br_run._r.append(OxmlElement("w:br"))
            continue
        if ct != "text":
            continue

        text = child.get("text", "")
        href = None
        bold = italic = strike = code = False
        for mark in child.get("marks", []) or []:
            mt = mark.get("type")
            if mt == "bold":
                bold = True
            elif mt == "italic":
                italic = True
            elif mt == "strike":
                strike = True
            elif mt == "code":
                code = True
            elif mt == "link":
                href = mark.get("attrs", {}).get("href", "")

        if href:
            _add_hyperlink(paragraph, text, href)
        else:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            if strike:
                run.font.strike = True
            if code:
                run.font.name = "Courier New"
                run.font.size = Pt(10)


def _build_docx(title: str, doc: dict) -> bytes:
    from docx import Document as DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    d = DocxDocument()
    _set_doc_font(d)
    if title:
        d.add_heading(title, level=0)

    for node in doc.get("content", []) or []:
        t = node.get("type")

        if t == "heading":
            level = min(max(int(node.get("attrs", {}).get("level", 1)), 1), 6)
            p = d.add_heading("", level=level)
            _run_inline(p, node)

        elif t == "paragraph":
            p = d.add_paragraph()
            _run_inline(p, node)

        elif t in ("bulletList", "orderedList"):
            style = "List Bullet" if t == "bulletList" else "List Number"
            for item in node.get("content", []) or []:
                for sub in item.get("content", []) or []:
                    p = d.add_paragraph(style=style)
                    _run_inline(p, sub)

        elif t == "blockquote":
            for inner in node.get("content", []) or []:
                try:
                    p = d.add_paragraph(style="Quote")
                except KeyError:
                    p = d.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                _run_inline(p, inner)

        elif t == "codeBlock":
            code_text = _node_text(node)
            p = d.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

        elif t == "horizontalRule":
            p = d.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif t == "table":
            rows = node.get("content", []) or []
            if not rows:
                continue
            col_count = max(len(r.get("content", []) or []) for r in rows)
            if col_count == 0:
                continue
            tbl = d.add_table(rows=len(rows), cols=col_count)
            tbl.style = "Table Grid"
            for ri, row_node in enumerate(rows):
                cells = row_node.get("content", []) or []
                for ci, cell_node in enumerate(cells):
                    if ci >= col_count:
                        break
                    cell = tbl.cell(ri, ci)
                    cell.paragraphs[0].clear()
                    cell_paras = cell_node.get("content", []) or []
                    for pi, para_node in enumerate(cell_paras):
                        p = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
                        _run_inline(p, para_node)
                        if pi < len(cell_paras) - 1:
                            cell.add_paragraph()

        elif t == "image":
            data = _decode_data_url(node.get("attrs", {}).get("src", ""))
            if data:
                try:
                    d.add_picture(io.BytesIO(data), width=Inches(5.5))
                except Exception:  # noqa: BLE001 — 图片损坏不阻断导出
                    logger.warning("case docx 跳过无法解析的图片")

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ── PDF 生成（reportlab + CJK 字体） ────────────────────────────────────────
_CJK_FONT_NAME = "STSong-Light"
_CJK_FONT_REGISTERED = False


def _ensure_cjk_font() -> str | None:
    """注册 reportlab 内置 CID 中文字体，返回字体名；失败返回 None（回退默认）。"""
    global _CJK_FONT_REGISTERED
    if _CJK_FONT_REGISTERED:
        return _CJK_FONT_NAME
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    try:
        pdfmetrics.registerFont(UnicodeCIDFont(_CJK_FONT_NAME))
        _CJK_FONT_REGISTERED = True
        return _CJK_FONT_NAME
    except Exception:  # noqa: BLE001
        logger.warning("case pdf 注册 CID 中文字体失败，中文可能显示异常")
        return None


def _pdf_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_markup(node: dict) -> str:
    """把内联节点转成 reportlab Paragraph 的迷你标记。
    支持 bold / italic / strike / code marks、link mark、hardBreak。
    """
    out: list[str] = []
    for child in node.get("content", []) or []:
        ct = child.get("type")
        if ct == "hardBreak":
            out.append("<br/>")
            continue
        if ct != "text":
            continue
        text = _pdf_escape(child.get("text", ""))
        href = None
        bold = italic = strike = code = False
        for mark in child.get("marks", []) or []:
            mt = mark.get("type")
            if mt == "bold":
                bold = True
            elif mt == "italic":
                italic = True
            elif mt == "strike":
                strike = True
            elif mt == "code":
                code = True
            elif mt == "link":
                href = mark.get("attrs", {}).get("href")
        if code:
            text = f'<font name="Courier">{text}</font>'
        if bold:
            text = f"<b>{text}</b>"
        if italic:
            text = f"<i>{text}</i>"
        if strike:
            text = f"<strike>{text}</strike>"
        if href:
            text = f'<a href="{_pdf_escape(href)}">{text}</a>'
        out.append(text)
    return "".join(out)


def _build_pdf(title: str, doc: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import (
        HRFlowable,
        Image,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        TableStyle,
    )
    from reportlab.platypus import (
        Table as PdfTable,
    )

    font = _ensure_cjk_font()
    styles = getSampleStyleSheet()
    base_font = font or "Helvetica"

    def _style(name: str, **kw) -> ParagraphStyle:
        return ParagraphStyle(name, parent=styles["Normal"], fontName=base_font, **kw)

    body_style = _style("CaseBody", fontSize=11, leading=16)
    title_style = _style("CaseTitle", fontSize=20, leading=26, spaceAfter=12)
    quote_style = _style("CaseQuote", fontSize=11, leading=16, leftIndent=24, textColor=colors.HexColor("#555555"))
    code_style = _style("CaseCode", fontName="Courier", fontSize=9, leading=14, leftIndent=12)
    heading_styles = {
        i: _style(
            f"CaseH{i}", fontSize=max(18 - i * 2, 12), leading=20, spaceBefore=8, spaceAfter=4
        )
        for i in range(1, 7)
    }

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4)
    flow: list = []
    if title:
        flow.append(Paragraph(_pdf_escape(title), title_style))

    for node in doc.get("content", []) or []:
        t = node.get("type")

        if t == "heading":
            level = min(max(int(node.get("attrs", {}).get("level", 1)), 1), 6)
            flow.append(Paragraph(_inline_markup(node) or "&nbsp;", heading_styles[level]))

        elif t == "paragraph":
            flow.append(Paragraph(_inline_markup(node) or "&nbsp;", body_style))

        elif t in ("bulletList", "orderedList"):
            items = []
            for item in node.get("content", []) or []:
                for sub in item.get("content", []) or []:
                    items.append(ListItem(Paragraph(_inline_markup(sub) or "&nbsp;", body_style)))
            if items:
                flow.append(
                    ListFlowable(items, bulletType="bullet" if t == "bulletList" else "1")
                )

        elif t == "blockquote":
            for inner in node.get("content", []) or []:
                flow.append(Paragraph(_inline_markup(inner) or "&nbsp;", quote_style))

        elif t == "codeBlock":
            code_text = _pdf_escape(_node_text(node)).replace("\n", "<br/>")
            flow.append(Paragraph(code_text or "&nbsp;", code_style))

        elif t == "horizontalRule":
            flow.append(HRFlowable(width="100%", thickness=1, color=colors.grey, spaceAfter=6))

        elif t == "table":
            rows = node.get("content", []) or []
            if not rows:
                continue
            table_data = []
            for row_node in rows:
                row = []
                for cell_node in row_node.get("content", []) or []:
                    cell_markup = " ".join(
                        _inline_markup(p) for p in cell_node.get("content", []) or []
                    ) or " "
                    row.append(Paragraph(cell_markup, body_style))
                if row:
                    table_data.append(row)
            if table_data:
                tbl = PdfTable(table_data, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.9)),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("FONTNAME", (0, 0), (-1, -1), base_font),
                ]))
                flow.append(tbl)
                flow.append(Spacer(1, 6))

        elif t == "image":
            data = _decode_data_url(node.get("attrs", {}).get("src", ""))
            if data:
                try:
                    img = Image(io.BytesIO(data))
                    max_w = 440
                    if img.drawWidth > max_w:
                        ratio = max_w / img.drawWidth
                        img.drawWidth = max_w
                        img.drawHeight = img.drawHeight * ratio
                    flow.append(img)
                    flow.append(Spacer(1, 6))
                except Exception:  # noqa: BLE001
                    logger.warning("case pdf 跳过无法解析的图片")

    if not flow:
        flow.append(Paragraph("&nbsp;", body_style))
    pdf.build(flow)
    return buf.getvalue()


# ── 对外入口 ───────────────────────────────────────────────────────────────
def export_case(
    *, title: str, fmt: str, content_json: dict, content_html: str | None = None
) -> tuple[bytes, str, str]:
    """把 Tiptap 文档导出为 docx/pdf。

    返回 (data, mime_type, plain_text)。fmt ∈ {"docx","pdf"}。
    """
    doc = content_json if isinstance(content_json, dict) else {}
    plain_text = _extract_plain_text(doc)
    if title:
        plain_text = f"{title}\n{plain_text}".strip()

    if fmt == "docx":
        return _build_docx(title, doc), DOCX_MIME, plain_text
    if fmt == "pdf":
        return _build_pdf(title, doc), PDF_MIME, plain_text
    raise ValueError(f"不支持的导出格式: {fmt!r}")
