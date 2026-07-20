"""对话导出服务：MessagePublic 列表 → .docx / .md 字节流。"""

import io
import re
import uuid
from datetime import date
from typing import Literal

from lxml import etree

from app.schemas.chat import MessagePublic

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MD_MIME = "text/markdown; charset=utf-8"
PDF_MIME = "application/pdf"

_MAX_FILENAME_BASE = 50

# 微软雅黑：Windows Vista+ 及所有 Office 版本原生支持的中文字体
_DOCX_FONT = "Microsoft YaHei"


def _set_doc_font(doc: object, font_name: str = _DOCX_FONT) -> None:
    """设置 docx 文档默认字体（ASCII + 东亚字体槽），使中文正确渲染。"""
    from docx.oxml.ns import qn

    rPr = doc.styles["Normal"].element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _safe_filename(title: str) -> str:
    """生成安全的文件名基础部分（无扩展名）。"""
    name = re.sub(r'[\\/:*?"<>|\r\n]', "_", title).strip()
    if not name:
        name = "conversation"
    if len(name) > _MAX_FILENAME_BASE:
        name = name[:_MAX_FILENAME_BASE]
    return name


def build_export(
    title: str,
    messages: list[MessagePublic],
    fmt: Literal["docx", "md", "pdf"],
    *,
    message_ids: list[uuid.UUID] | None = None,
) -> tuple[bytes, str, str]:
    """返回 (bytes, mime_type, filename)。"""
    selected = [m for m in messages if message_ids is None or m.id in set(message_ids)]
    date_str = date.today().isoformat()
    base = _safe_filename(title)
    if fmt == "md":
        data = _build_markdown(title, selected, date_str)
        return data, MD_MIME, f"{base}_{date_str}.md"
    if fmt == "pdf":
        data = _build_pdf(title, selected, date_str)
        return data, PDF_MIME, f"{base}_{date_str}.pdf"
    data = _build_docx(title, selected, date_str)
    return data, DOCX_MIME, f"{base}_{date_str}.docx"


# ── Markdown ────────────────────────────────────────────────────────────────

def _build_markdown(title: str, messages: list[MessagePublic], date_str: str) -> bytes:
    lines: list[str] = [f"# {title}", "", f"*导出时间：{date_str}*", ""]
    for msg in messages:
        if msg.role == "user":
            lines += ["---", "", "**提问**", "", msg.content, ""]
        else:
            lines += ["**回答**", "", msg.content, ""]
            sources: list = getattr(msg, "sources", []) or []
            if sources:
                titles = " · ".join(
                    s.get("title", "") if isinstance(s, dict) else getattr(s, "title", "")
                    for s in sources
                )
                if titles:
                    lines += [f"> 来源：{titles}", ""]
    lines.append("---")
    return "\n".join(lines).encode("utf-8")


# ── DOCX ────────────────────────────────────────────────────────────────────

_WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Inline Markdown patterns: order matters — bold+italic first, then bold, then italic
_INLINE_RE = re.compile(
    r"(\*\*\*|___)(.*?)\1"          # bold+italic ***x*** or ___x___
    r"|(\*\*|__)(.*?)\3"            # bold **x** or __x__
    r"|(\*|_)(.*?)\5"               # italic *x* or _x_
    r"|`([^`]+)`",                  # inline code `x`
    re.DOTALL,
)


def _add_inline_runs(para: object, text: str, base_size: object = None) -> None:
    """Parse inline Markdown and add styled runs to *para*."""
    from docx.shared import RGBColor

    pos = 0
    for m in _INLINE_RE.finditer(text):
        # plain text before this match
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            if base_size:
                r.font.size = base_size
        g1 = m.group(1)
        g2 = m.group(2)
        g3 = m.group(3)
        g4 = m.group(4)
        g5 = m.group(5)
        g6 = m.group(6)
        g7 = m.group(7)
        if g1:  # bold+italic
            r = para.add_run(g2)
            r.bold = True
            r.italic = True
        elif g3:  # bold
            r = para.add_run(g4)
            r.bold = True
        elif g5:  # italic
            r = para.add_run(g6)
            r.italic = True
        else:  # inline code
            r = para.add_run(g7)
            r.font.name = "Courier New"
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        if base_size:
            r.font.size = base_size
        pos = m.end()
    # trailing plain text
    if pos < len(text):
        r = para.add_run(text[pos:])
        if base_size:
            r.font.size = base_size


def _add_hr(doc: object) -> None:
    """Add a horizontal rule paragraph."""
    sep = doc.add_paragraph()
    pPr = sep._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, f"{{{_WNS}}}pBdr")
    bottom = etree.SubElement(pBdr, f"{{{_WNS}}}bottom")
    bottom.set(f"{{{_WNS}}}val", "single")
    bottom.set(f"{{{_WNS}}}sz", "4")
    bottom.set(f"{{{_WNS}}}space", "1")
    bottom.set(f"{{{_WNS}}}color", "CCCCCC")


def _render_markdown_to_doc(doc: object, md: str) -> None:
    """Render *md* Markdown text into *doc* (python-docx Document)."""
    from docx.shared import Pt, RGBColor

    lines = md.splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    def flush_code(lines_buf: list[str]) -> None:
        joined = "\n".join(lines_buf)
        p = doc.add_paragraph()
        r = p.add_run(joined)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        # light grey paragraph shading via XML
        pPr = p._p.get_or_add_pPr()
        shd = etree.SubElement(pPr, f"{{{_WNS}}}shd")
        shd.set(f"{{{_WNS}}}val", "clear")
        shd.set(f"{{{_WNS}}}color", "auto")
        shd.set(f"{{{_WNS}}}fill", "F3F4F6")

    while i < len(lines):
        line = lines[i]

        # ── fenced code block ─────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                flush_code(code_lines)
                in_code_block = False
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── headings ──────────────────────────────────────────────────────
        hm = re.match(r"^(#{1,6})\s+(.*)", line)
        if hm:
            level = len(hm.group(1))
            p = doc.add_heading("", level=min(level, 4))
            _add_inline_runs(p, hm.group(2).strip())
            i += 1
            continue

        # ── horizontal rule ───────────────────────────────────────────────
        if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", line):
            _add_hr(doc)
            i += 1
            continue

        # ── blockquote ────────────────────────────────────────────────────
        if line.startswith("> "):
            style = "Quote" if _style_exists(doc, "Quote") else None
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            _add_inline_runs(p, line[2:], Pt(10))
            i += 1
            continue

        # ── bullet list ───────────────────────────────────────────────────
        bm = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if bm:
            indent = len(bm.group(1)) // 2
            style = "List Bullet" if indent == 0 else "List Bullet 2"
            p = doc.add_paragraph(style=style) if _style_exists(doc, style) else doc.add_paragraph()
            _add_inline_runs(p, bm.group(2))
            i += 1
            continue

        # ── numbered list ─────────────────────────────────────────────────
        nm = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
        if nm:
            indent = len(nm.group(1)) // 2
            style = "List Number" if indent == 0 else "List Number 2"
            p = doc.add_paragraph(style=style) if _style_exists(doc, style) else doc.add_paragraph()
            _add_inline_runs(p, nm.group(2))
            i += 1
            continue

        # ── blank line → paragraph break ─────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── normal paragraph (may span multiple non-blank lines) ──────────
        chunk_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") \
                and not lines[i].startswith("```") and not lines[i].startswith("> ") \
                and not re.match(r"^(\s*)[-*+]\s+", lines[i]) \
                and not re.match(r"^(\s*)\d+[.)]\s+", lines[i]):
            chunk_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(chunk_lines))

    # flush unclosed code block
    if in_code_block and code_lines:
        flush_code(code_lines)


def _style_exists(doc: object, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def _build_docx(title: str, messages: list[MessagePublic], date_str: str) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor

    d = DocxDocument()
    _set_doc_font(d)

    # 标题
    heading = d.add_heading(title, level=1)
    heading.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    # 元数据行
    meta = d.add_paragraph(f"导出时间：{date_str}")
    meta.alignment = 1
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    for msg in messages:
        # 角色标签
        if msg.role == "user":
            label, color = "「提问」", RGBColor(0x1D, 0x4E, 0xD8)
        else:
            label, color = "「回答」", RGBColor(0x06, 0x5F, 0x46)

        label_para = d.add_paragraph()
        label_run = label_para.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = color

        # 正文：用户消息直接输出，AI 消息解析 Markdown
        if msg.role == "user":
            for line in msg.content.splitlines():
                d.add_paragraph(line)
        else:
            _render_markdown_to_doc(d, msg.content)

        # 来源
        sources: list = getattr(msg, "sources", []) or []
        if sources and msg.role == "assistant":
            titles = " · ".join(
                s.get("title", "") if isinstance(s, dict) else getattr(s, "title", "")
                for s in sources
            )
            if titles:
                src_para = d.add_paragraph(f"来源：{titles}")
                for run in src_para.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        _add_hr(d)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ── PDF ─────────────────────────────────────────────────────────────────────

_CJK_FONT_NAME = "STSong-Light"
_CJK_FONT_REGISTERED = False


def _ensure_cjk_font() -> str:
    global _CJK_FONT_REGISTERED
    if _CJK_FONT_REGISTERED:
        return _CJK_FONT_NAME
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    try:
        pdfmetrics.registerFont(UnicodeCIDFont(_CJK_FONT_NAME))
        _CJK_FONT_REGISTERED = True
    except Exception:  # noqa: BLE001
        return "Helvetica"
    return _CJK_FONT_NAME


def _pdf_esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# Inline Markdown → reportlab XML markup
_RL_INLINE_RE = re.compile(
    r"(\*\*\*|___)(.*?)\1"
    r"|(\*\*|__)(.*?)\3"
    r"|(\*|_)(.*?)\5"
    r"|`([^`]+)`",
    re.DOTALL,
)


def _md_inline_to_rl(text: str) -> str:
    """Convert inline Markdown to reportlab Paragraph XML."""
    result: list[str] = []
    pos = 0
    for m in _RL_INLINE_RE.finditer(text):
        if m.start() > pos:
            result.append(_pdf_esc(text[pos:m.start()]))
        g1, g2, g3, g4, g5, g6, g7 = (
            m.group(1), m.group(2), m.group(3), m.group(4),
            m.group(5), m.group(6), m.group(7),
        )
        if g1:
            result.append(f"<b><i>{_pdf_esc(g2)}</i></b>")
        elif g3:
            result.append(f"<b>{_pdf_esc(g4)}</b>")
        elif g5:
            result.append(f"<i>{_pdf_esc(g6)}</i>")
        else:
            result.append(f'<font face="Courier" size="9">{_pdf_esc(g7)}</font>')
        pos = m.end()
    if pos < len(text):
        result.append(_pdf_esc(text[pos:]))
    return "".join(result)


def _build_pdf(title: str, messages: list[MessagePublic], date_str: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        HRFlowable,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    font = _ensure_cjk_font()
    styles = getSampleStyleSheet()

    def st(name: str, **kw: object) -> ParagraphStyle:
        return ParagraphStyle(name, parent=styles["Normal"],
                              fontName=kw.pop("fontName", font), **kw)

    title_st = st("ExTitle", fontSize=18, leading=24, spaceAfter=4, alignment=1)
    meta_st = st("ExMeta", fontSize=9, leading=12, textColor=colors.grey,
                 spaceAfter=10, alignment=1)
    label_q_st = st("ExLabelQ", fontSize=11, leading=14, textColor=colors.HexColor("#1D4ED8"),
                    spaceBefore=10, spaceAfter=4)
    label_a_st = st("ExLabelA", fontSize=11, leading=14, textColor=colors.HexColor("#065F46"),
                    spaceBefore=10, spaceAfter=4)
    body_st = st("ExBody", fontSize=11, leading=16, spaceAfter=4)
    code_st = st("ExCode", fontSize=9, leading=13, fontName="Courier",
                 backColor=colors.HexColor("#F3F4F6"), leftIndent=8, spaceAfter=6)
    src_st = st("ExSrc", fontSize=9, leading=12, textColor=colors.grey, spaceAfter=4)
    h_styles = {
        i: st(f"ExH{i}", fontSize=max(16 - i * 2, 10), leading=max(20 - i * 2, 14),
              spaceBefore=8, spaceAfter=4)
        for i in range(1, 7)
    }
    bq_st = st("ExBQ", fontSize=10, leading=14, leftIndent=12,
               textColor=colors.HexColor("#555555"), spaceAfter=4)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20 * mm, rightMargin=20 * mm,
                            topMargin=20 * mm, bottomMargin=20 * mm)
    flow: list = [
        Paragraph(_pdf_esc(title), title_st),
        Paragraph(f"导出时间：{date_str}", meta_st),
        HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey, spaceAfter=8),
    ]

    for msg in messages:
        label = "「提问」" if msg.role == "user" else "「回答」"
        flow.append(Paragraph(label, label_q_st if msg.role == "user" else label_a_st))

        if msg.role == "user":
            # User messages: plain text, preserve line breaks
            for line in msg.content.splitlines():
                if line.strip():
                    flow.append(Paragraph(_pdf_esc(line), body_st))
        else:
            # AI messages: parse Markdown
            lines = msg.content.splitlines()
            i = 0
            in_code = False
            code_lines: list[str] = []

            def flush_code(buf_lines: list[str]) -> None:
                flow.append(Paragraph(_pdf_esc("\n".join(buf_lines)), code_st))

            while i < len(lines):
                line = lines[i]
                if line.startswith("```"):
                    if not in_code:
                        in_code = True
                        code_lines = []
                    else:
                        flush_code(code_lines)
                        in_code = False
                        code_lines = []
                    i += 1
                    continue
                if in_code:
                    code_lines.append(line)
                    i += 1
                    continue
                hm = re.match(r"^(#{1,6})\s+(.*)", line)
                if hm:
                    lvl = min(len(hm.group(1)), 6)
                    flow.append(Paragraph(_md_inline_to_rl(hm.group(2).strip()), h_styles[lvl]))
                    i += 1
                    continue
                if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", line):
                    flow.append(HRFlowable(width="100%", thickness=0.5,
                                           color=colors.lightgrey, spaceAfter=4))
                    i += 1
                    continue
                if line.startswith("> "):
                    flow.append(Paragraph(_md_inline_to_rl(line[2:]), bq_st))
                    i += 1
                    continue
                bm = re.match(r"^(\s*)[-*+]\s+(.*)", line)
                if bm:
                    items = [ListItem(Paragraph(_md_inline_to_rl(bm.group(2)), body_st))]
                    i += 1
                    while i < len(lines):
                        nm2 = re.match(r"^(\s*)[-*+]\s+(.*)", lines[i])
                        if nm2:
                            items.append(ListItem(
                                Paragraph(_md_inline_to_rl(nm2.group(2)), body_st)))
                            i += 1
                        else:
                            break
                    flow.append(ListFlowable(items, bulletType="bullet"))
                    continue
                nm = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
                if nm:
                    items = [ListItem(Paragraph(_md_inline_to_rl(nm.group(2)), body_st))]
                    i += 1
                    while i < len(lines):
                        nm2 = re.match(r"^(\s*)\d+[.)]\s+(.*)", lines[i])
                        if nm2:
                            items.append(ListItem(
                                Paragraph(_md_inline_to_rl(nm2.group(2)), body_st)))
                            i += 1
                        else:
                            break
                    flow.append(ListFlowable(items, bulletType="1"))
                    continue
                if line.strip() == "":
                    i += 1
                    continue
                # normal paragraph — collect until blank/block
                chunk = [line]
                i += 1
                while i < len(lines) and lines[i].strip() \
                        and not lines[i].startswith("#") \
                        and not lines[i].startswith("```") \
                        and not lines[i].startswith("> ") \
                        and not re.match(r"^(\s*)[-*+]\s+", lines[i]) \
                        and not re.match(r"^(\s*)\d+[.)]\s+", lines[i]):
                    chunk.append(lines[i])
                    i += 1
                flow.append(Paragraph(_md_inline_to_rl(" ".join(chunk)), body_st))

            if in_code and code_lines:
                flush_code(code_lines)

        # Sources
        sources: list = getattr(msg, "sources", []) or []
        if sources and msg.role == "assistant":
            titles = " · ".join(
                s.get("title", "") if isinstance(s, dict) else getattr(s, "title", "")
                for s in sources
            )
            if titles:
                flow.append(Paragraph(f"来源：{_pdf_esc(titles)}", src_st))

        flow.append(HRFlowable(width="100%", thickness=0.5,
                               color=colors.lightgrey, spaceAfter=6))
        flow.append(Spacer(1, 2))

    doc.build(flow)
    return buf.getvalue()
